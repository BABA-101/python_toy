# 수료증 레이아웃을 워드 문서로 만들기
# docx 라이브러리 : https://python-docx.readthedocs.io/en/latest/
import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc=docx.Document(r'project-12\수료증양식.docx')

style = doc.styles['Normal']
style.font.name='나눔고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
style.font.size=docx.shared.Pt(12)

#문단 생성, 글시 입력 후 폰트와 글씨크기 설정
para = doc.add_paragraph()
run=para.add_run('\n')
run=para.add_run(' 제 2020-0001호\n')
run.font.name = '나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size=docx.shared.Pt(20)

para = doc.add_paragraph()
run=para.add_run('\n')
run=para.add_run('수료증')
run.font.name = '나눔고딕'
run.bold=True
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size=docx.shared.Pt(40)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para=doc.add_paragraph()
run=para.add_run('\n')
run=para.add_run(' 성 명: 홍길동\n')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run=para.add_run(' 생 년 월 일: 1990.01.01\n')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run=para.add_run(' 교 육 과 정: 파이썬과 40개의 작품들\n')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run=para.add_run(' 교 육 날 짜: 2022.08.05~2022.10.05\n')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)

para=doc.add_paragraph()
run=para.add_run('\n\n')
run=para.add_run(' 위 사람은 파이썬과 40개의 작품들 교육과정을')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
run=para.add_run(' 이수하였으므로 이 증서를 수여 합니다.')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)

para=doc.add_paragraph()
run=para.add_run('\n\n\n')
run=para.add_run(' 2022.10.27')
run.font.name='나눔고딕'
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

para=doc.add_paragraph()
run=para.add_run('\n\n\n')
run=para.add_run(' 파이썬교육기관장')
run.font.name='나눔고딕'
run.bold=True
run._element.rPr.rFonts.set(qn('w:eastAsia'),'나눔고딕')
run.font.size = docx.shared.Pt(20)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(r'project-12\수료증결과.docx')
